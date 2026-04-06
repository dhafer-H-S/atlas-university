"""
build_report.py  –  generates report/CacheLocalityLab_Report.docx
Istanbul Atlas University | Computer Systems | Mid-term Project | Team 11
"""

import os, math, csv
from pathlib import Path

# ── third-party ─────────────────────────────────────────────────────────────
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text  import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns   import qn
from docx.oxml      import OxmlElement

# ── paths ────────────────────────────────────────────────────────────────────
BASE    = Path(__file__).resolve().parent.parent
RESULTS = BASE / "results"
FIGS    = BASE / "results" / "figures"
FIGS.mkdir(parents=True, exist_ok=True)
OUT     = BASE / "report" / "CacheLocalityLab_Report.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

# ── colour palette (dark-academic) ───────────────────────────────────────────
BG   = "#0f0f1a"
BG2  = "#1a1a30"
FG   = "#e8e8ff"
TEAL = "#2dd4bf"
BLUE = "#60a5fa"
AMB  = "#fbbf24"
RED  = "#f87171"
GRN  = "#34d399"
GRID = "#2a2a45"

plt.rcParams.update({
    "figure.facecolor": BG,
    "axes.facecolor":   BG2,
    "axes.edgecolor":   "#444466",
    "axes.labelcolor":  FG,
    "xtick.color":      FG,
    "ytick.color":      FG,
    "text.color":       FG,
    "grid.color":       GRID,
    "grid.linestyle":   "--",
    "grid.alpha":       0.6,
    "font.family":      "sans-serif",
    "figure.dpi":       150,
})

# ═══════════════════════════════════════════════════════════════════════════
# 1. CHART GENERATION
# ═══════════════════════════════════════════════════════════════════════════

def load_csv(name):
    rows = []
    with open(RESULTS / name) as f:
        reader = csv.DictReader(f)
        for r in reader:
            rows.append({k: float(v) if v.replace(".", "", 1).lstrip("-").isdigit() else v
                         for k, v in r.items()})
    return rows

# ── Figure 1: Stride ─────────────────────────────────────────────────────────
def fig_stride():
    data   = load_csv("stride.csv")
    strides= [r["stride_bytes"] for r in data]
    ns     = [r["ns_per_access"] for r in data]
    colors = [RED if n > 2 else TEAL for n in ns]

    fig, axes = plt.subplots(1, 2, figsize=(12, 4.5), facecolor=BG)
    ax = axes[0]
    bars = ax.bar(range(len(strides)), ns, color=colors, edgecolor="#0f0f1a", linewidth=0.8)
    ax.set_xticks(range(len(strides)))
    ax.set_xticklabels([f"{int(s)}B" for s in strides], rotation=35, ha="right", fontsize=8)
    ax.set_ylabel("ns / access", fontsize=10)
    ax.set_title("Figure 1.1  Stride Access Latency", fontsize=11, color=FG, pad=8)
    ax.axhline(2.0, color=AMB, linestyle="--", linewidth=1.2, label="2 ns threshold")
    ax.axvline(4.5, color="#f97316", linestyle=":", linewidth=1.5, label="64 B boundary")
    ax.legend(fontsize=8)
    ax.grid(axis="y")
    for bar, val in zip(bars, ns):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05,
                f"{val:.2f}", ha="center", va="bottom", fontsize=7, color=FG)

    # table panel
    ax2 = axes[1]
    ax2.axis("off")
    col_labels = ["Stride (B)", "ns/access", "Cache status"]
    table_data = []
    for r in data:
        status = "L1/L2 hit" if r["ns_per_access"] < 2 else "Cache miss"
        table_data.append([f"{int(r['stride_bytes'])}", f"{r['ns_per_access']:.3f}", status])
    t = ax2.table(cellText=table_data, colLabels=col_labels,
                  loc="center", cellLoc="center")
    t.auto_set_font_size(False)
    t.set_fontsize(8.5)
    t.scale(1.1, 1.55)
    for (row, col), cell in t.get_celld().items():
        cell.set_facecolor(BG2 if row > 0 else "#1e3a5f")
        cell.set_edgecolor("#444466")
        cell.set_text_props(color=FG)
    ax2.set_title("Table 1.1  Raw stride measurements", fontsize=10, color=FG, pad=8)

    plt.tight_layout(pad=1.5)
    path = FIGS / "report_stride.png"
    plt.savefig(path, bbox_inches="tight", facecolor=BG)
    plt.close()
    print(f"  ✓ {path.name}")
    return path

# ── Figure 2: Working Set ─────────────────────────────────────────────────────
def fig_workingset():
    data  = load_csv("workingset.csv")
    kb    = [r["array_kb"] for r in data]
    ns    = [r["ns_per_access"] for r in data]

    fig, ax = plt.subplots(figsize=(10, 5), facecolor=BG)
    ax.plot(range(len(kb)), ns, color=TEAL, linewidth=2.5, marker="o", markersize=5, label="ns/access")
    ax.set_xticks(range(len(kb)))
    ax.set_xticklabels([f"{int(k) if k < 1024 else int(k//1024)}{'' if k < 1024 else 'M'} KB"
                         if k < 1024 else f"{int(k/1024)} MB"
                         for k in kb], rotation=40, ha="right", fontsize=8)
    ax.set_ylabel("ns / access", fontsize=10)
    ax.set_title("Figure 2.1  Working Set Latency — Sequential Scan on Apple Silicon", fontsize=11, color=FG, pad=8)
    ax.axhline(1.05, color=AMB, linestyle="--", linewidth=1, alpha=0.7, label="~1 ns (UMA floor)")

    # shade cache regions
    regions = [(0, 5, "#0ea5e955", "L1 ≤32 KB"), (5, 8, "#10b98133", "L2 ≤256 KB"),
               (8, 14, "#f59e0b22", "L3 ≤6 MB"), (14, len(kb), "#dc262622", "DRAM")]
    for start, end, col, lbl in regions:
        ax.axvspan(start, min(end, len(kb)-1), facecolor=col, label=lbl)

    ax.legend(fontsize=8, loc="upper right")
    ax.grid(axis="y")

    annot = ("Apple Silicon UMA + System Level Cache\nhides all latency transitions for sequential access.\nAll measurements cluster at ~1 ns.")
    ax.annotate(annot, xy=(10, 1.01), xytext=(11, 1.025),
                fontsize=7.5, color=AMB,
                arrowprops=dict(arrowstyle="->", color=AMB, lw=1),
                bbox=dict(boxstyle="round,pad=0.3", fc="#1a1a30", ec=AMB, lw=1))

    plt.tight_layout(pad=1.5)
    path = FIGS / "report_workingset.png"
    plt.savefig(path, bbox_inches="tight", facecolor=BG)
    plt.close()
    print(f"  ✓ {path.name}")
    return path

# ── Figure 3: Matrix ─────────────────────────────────────────────────────────
def fig_matrix():
    data  = load_csv("matrix.csv")
    ns    = [r["N"] for r in data]
    row_v = [r["rowmajor_ns"] for r in data]
    col_v = [r["colmajor_ns"] for r in data]
    ratio = [r["slowdown_ratio"] for r in data]

    x = np.arange(len(ns))
    w = 0.35
    fig, (ax, ax2) = plt.subplots(1, 2, figsize=(12, 5), facecolor=BG)

    ax.bar(x - w/2, row_v, w, color=GRN,  label="Row-major",    edgecolor=BG, linewidth=0.8)
    ax.bar(x + w/2, col_v, w, color=RED,  label="Column-major", edgecolor=BG, linewidth=0.8)
    ax.set_xticks(x); ax.set_xticklabels([f"N={int(n)}" for n in ns], fontsize=9)
    ax.set_ylabel("ns / element", fontsize=10)
    ax.set_title("Figure 3.1  Absolute Traversal Latency", fontsize=11, color=FG, pad=8)
    ax.legend(fontsize=9); ax.grid(axis="y")

    ratio_colors = [AMB if r > 1.3 else TEAL for r in ratio]
    bars2 = ax2.bar(x, ratio, color=ratio_colors, edgecolor=BG, linewidth=0.8)
    ax2.axhline(8.0, color=RED, linestyle="--", linewidth=1.2, label="Theoretical max (8×)")
    ax2.axhline(1.0, color=TEAL, linestyle=":",  linewidth=1.0, label="Baseline 1×")
    ax2.set_xticks(x); ax2.set_xticklabels([f"N={int(n)}" for n in ns], fontsize=9)
    ax2.set_ylabel("Slowdown ratio (col / row)", fontsize=10)
    ax2.set_title("Figure 3.2  Column-major Slowdown Ratio", fontsize=11, color=FG, pad=8)
    ax2.legend(fontsize=8); ax2.grid(axis="y")
    for bar, val in zip(bars2, ratio):
        ax2.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.02,
                 f"{val:.2f}×", ha="center", va="bottom", fontsize=8, color=FG)

    plt.tight_layout(pad=1.5)
    path = FIGS / "report_matrix.png"
    plt.savefig(path, bbox_inches="tight", facecolor=BG)
    plt.close()
    print(f"  ✓ {path.name}")
    return path

# ── Figure 4: Linked List ─────────────────────────────────────────────────────
def fig_linkedlist():
    data = load_csv("linkedlist.csv")
    kb   = [r["list_kb"] for r in data]
    ns   = [r["ns_per_hop"] for r in data]

    fig, ax = plt.subplots(figsize=(10, 5), facecolor=BG)
    ax.plot(range(len(kb)), ns, color=AMB, linewidth=2.5, marker="o", markersize=5, label="Measured ns/hop")

    regions = [(0, 5, "#0ea5e955", "L1 ≤32 KB"), (5, 8, "#10b98133", "L2 ≤256 KB"),
               (8, 13, "#f59e0b22", "L3 ≤6 MB"), (13, len(kb), "#dc262622", "DRAM")]
    for start, end, col, lbl in regions:
        ax.axvspan(start, min(end, len(kb)-1), facecolor=col, label=lbl)

    ax.set_xticks(range(len(kb)))
    ax.set_xticklabels([f"{int(k)} KB" if k < 1024 else f"{int(k/1024)} MB"
                         for k in kb], rotation=40, ha="right", fontsize=8)
    ax.set_ylabel("ns / pointer hop", fontsize=10)
    ax.set_title("Figure 4.1  Pointer-Chasing Latency vs List Size", fontsize=11, color=FG, pad=8)
    ax.legend(fontsize=8, loc="upper left"); ax.grid(axis="y")

    ax.annotate("L1/L2: ~0.65 ns", xy=(3, 0.65), xytext=(4, 5),
                fontsize=8, color=GRN,
                arrowprops=dict(arrowstyle="->", color=GRN, lw=1))
    ax.annotate("DRAM: ~83 ns", xy=(len(kb)-1, ns[-1]), xytext=(len(kb)-4, 60),
                fontsize=8, color=RED,
                arrowprops=dict(arrowstyle="->", color=RED, lw=1))

    plt.tight_layout(pad=1.5)
    path = FIGS / "report_linkedlist.png"
    plt.savefig(path, bbox_inches="tight", facecolor=BG)
    plt.close()
    print(f"  ✓ {path.name}")
    return path

# ── Figure 5: Bandwidth ───────────────────────────────────────────────────────
def fig_bandwidth():
    data  = load_csv("bandwidth.csv")
    kb    = [r["array_kb"] for r in data]
    copy  = [r["copy_gb_per_s"] for r in data]
    sumv  = [r["sum_gb_per_s"] for r in data]

    fig, (ax, ax2) = plt.subplots(1, 2, figsize=(13, 5), facecolor=BG)

    ax.plot(range(len(kb)), copy, color=BLUE, linewidth=2.5, marker="o", markersize=4,  label="Copy (GB/s)")
    ax.plot(range(len(kb)), sumv, color=AMB,  linewidth=2.5, marker="s", markersize=4,  label="Sum (GB/s)")
    ax.set_xticks(range(0, len(kb), 2))
    ax.set_xticklabels([f"{int(k)} KB" if k < 1024 else f"{int(k/1024)} MB"
                         for k in kb[::2]], rotation=40, ha="right", fontsize=8)
    ax.set_ylabel("GB/s", fontsize=10)
    ax.set_title("Figure 5.1  Bandwidth vs Array Size", fontsize=11, color=FG, pad=8)
    ax.legend(fontsize=9); ax.grid(axis="y")

    # Peak per level grouped bar
    levels   = ["L1 (4–8 KB)", "L2 (~256 KB)", "L3 (~4 MB)", "DRAM (>64 MB)"]
    copy_bw  = [max(copy[:3]), copy[6], copy[10], np.mean(copy[-5:])]
    sum_bw   = [np.mean(sumv[:3]), sumv[6], sumv[10], np.mean(sumv[-5:])]
    x = np.arange(4); w = 0.35
    ax2.bar(x - w/2, copy_bw, w, color=BLUE, label="Copy", edgecolor=BG, linewidth=0.8)
    ax2.bar(x + w/2, sum_bw,  w, color=AMB,  label="Sum",  edgecolor=BG, linewidth=0.8)
    ax2.set_xticks(x); ax2.set_xticklabels(levels, fontsize=8.5)
    ax2.set_ylabel("GB/s", fontsize=10)
    ax2.set_title("Figure 5.2  Peak Bandwidth by Cache Level", fontsize=11, color=FG, pad=8)
    ax2.legend(fontsize=9); ax2.grid(axis="y")
    for i, (cv, sv) in enumerate(zip(copy_bw, sum_bw)):
        ax2.text(i - w/2, cv + 2, f"{cv:.0f}", ha="center", fontsize=8, color=FG)
        ax2.text(i + w/2, sv + 2, f"{sv:.0f}", ha="center", fontsize=8, color=FG)

    plt.tight_layout(pad=1.5)
    path = FIGS / "report_bandwidth.png"
    plt.savefig(path, bbox_inches="tight", facecolor=BG)
    plt.close()
    print(f"  ✓ {path.name}")
    return path

# ── Figure 6: Dashboard ───────────────────────────────────────────────────────
def fig_dashboard(stride_data, ws_data, matrix_data, bw_data):
    fig, axes = plt.subplots(2, 2, figsize=(14, 9), facecolor=BG)
    fig.suptitle("Figure 6.1  Results Dashboard — Cache & Locality Performance Lab",
                 fontsize=13, color=FG, y=0.98)

    # top-left: stride
    ax = axes[0][0]
    sb = [r["stride_bytes"] for r in stride_data]
    sn = [r["ns_per_access"] for r in stride_data]
    ax.bar(range(len(sb)), sn, color=[RED if v > 2 else TEAL for v in sn], linewidth=0)
    ax.set_xticks(range(len(sb)))
    ax.set_xticklabels([f"{int(s)}B" for s in sb], fontsize=7, rotation=30, ha="right")
    ax.set_title("Exp 1: Stride", color=FG); ax.grid(axis="y"); ax.set_ylabel("ns/access", fontsize=9)

    # top-right: working set
    ax = axes[0][1]
    wkb = [r["array_kb"] for r in ws_data]
    wns = [r["ns_per_access"] for r in ws_data]
    ax.plot(range(len(wkb)), wns, color=TEAL, linewidth=2, marker="o", markersize=4)
    ax.set_xticks(range(0, len(wkb), 3))
    ax.set_xticklabels([f"{int(k)}K" if k < 1024 else f"{int(k/1024)}M"
                         for k in wkb[::3]], fontsize=7, rotation=30, ha="right")
    ax.set_title("Exp 2: Working Set", color=FG); ax.grid(axis="y"); ax.set_ylabel("ns/access", fontsize=9)
    ax.axhline(1.0, color=AMB, linestyle="--", linewidth=1, alpha=0.7)

    # bottom-left: matrix slowdown
    ax = axes[1][0]
    mns = [r["N"] for r in matrix_data]
    mr  = [r["slowdown_ratio"] for r in matrix_data]
    ax.bar(range(len(mns)), mr, color=[AMB if v > 1.3 else TEAL for v in mr], linewidth=0)
    ax.axhline(8.0, color=RED, linestyle="--", linewidth=1, alpha=0.8)
    ax.set_xticks(range(len(mns)))
    ax.set_xticklabels([f"N={int(n)}" for n in mns], fontsize=8)
    ax.set_title("Exp 3: Matrix Slowdown", color=FG); ax.grid(axis="y"); ax.set_ylabel("ratio", fontsize=9)

    # bottom-right: bandwidth
    ax = axes[1][1]
    bkb  = [r["array_kb"] for r in bw_data]
    bcopy= [r["copy_gb_per_s"] for r in bw_data]
    bsum = [r["sum_gb_per_s"] for r in bw_data]
    ax.plot(range(len(bkb)), bcopy, color=BLUE, linewidth=2, label="Copy")
    ax.plot(range(len(bkb)), bsum,  color=AMB,  linewidth=2, label="Sum")
    ax.set_xticks(range(0, len(bkb), 3))
    ax.set_xticklabels([f"{int(k)}K" if k < 1024 else f"{int(k/1024)}M"
                         for k in bkb[::3]], fontsize=7, rotation=30, ha="right")
    ax.set_title("Exp 5: Bandwidth", color=FG); ax.grid(axis="y")
    ax.set_ylabel("GB/s", fontsize=9); ax.legend(fontsize=8)

    plt.tight_layout(rect=[0, 0, 1, 0.97], pad=1.5)
    path = FIGS / "report_dashboard.png"
    plt.savefig(path, bbox_inches="tight", facecolor=BG)
    plt.close()
    print(f"  ✓ {path.name}")
    return path

# ═══════════════════════════════════════════════════════════════════════════
# 2. DOCX HELPERS
# ═══════════════════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=10, color="000000", align="left"):
    para = cell.paragraphs[0]
    para.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                      "center": WD_ALIGN_PARAGRAPH.CENTER,
                      "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    run  = para.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return para

def add_heading(doc, text, level):
    """Add heading with custom styling that matches template requirements."""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(11, 45, 82)  # dark navy
        if level == 1:
            run.font.size = Pt(16)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    return h

def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name      = "Times New Roman"
    run.font.size      = Pt(12)
    run.font.italic    = italic
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    # 1.5 line spacing
    pf = p.paragraph_format
    pf.line_spacing = Pt(18)
    pf.space_after  = Pt(6)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run  = p.add_run(text)
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x55, 0x66)
    p.paragraph_format.space_after = Pt(10)
    return p

def add_figure(doc, img_path, width_in=6.2, caption=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_in))
    if caption:
        add_caption(doc, caption)

def add_code_block(doc, code_text):
    """Add a monospaced code-style paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    pf = p.paragraph_format
    pf.left_indent  = Cm(1.2)
    pf.right_indent = Cm(1.2)
    pf.space_before = Pt(6)
    pf.space_after  = Pt(6)
    # light grey shading behind code
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "EFF3F8")
    pPr.append(shd)
    return p

def add_data_table(doc, headers, rows, caption=""):
    """Styled data table with atlas-blue header row."""
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.bold = True; cr.font.size = Pt(10.5)
        cr.font.color.rgb = RGBColor(11, 45, 82)
        cr.font.name = "Times New Roman"

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_bg(cell, "0B2D52")
        cell_text(cell, h, bold=True, size=10, color="FFFFFF", align="center")

    # data rows
    for i, row in enumerate(rows):
        row_color = "F0F4F8" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            set_cell_bg(cell, row_color)
            cell_text(cell, str(val), size=10, color="111827", align="center")

    doc.add_paragraph()  # spacing after table

def add_spacer(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

# ═══════════════════════════════════════════════════════════════════════════
# 3. DOCUMENT ASSEMBLY
# ═══════════════════════════════════════════════════════════════════════════

def build_doc(figs):
    doc = Document()

    # ── page margins (≈ template: 2.5 cm sides, 3 cm top/bottom) ────────
    for section in doc.sections:
        section.top_margin    = Cm(3.0)
        section.bottom_margin = Cm(3.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── default font (template: Times New Roman 12pt) ────────────────────
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    # ════════════════════════════════════════════════════
    # TITLE PAGE
    # ════════════════════════════════════════════════════
    add_spacer(doc, 4)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("Cache & Locality Performance Lab")
    tr.font.name = "Times New Roman"; tr.font.size = Pt(22); tr.bold = True
    tr.font.color.rgb = RGBColor(11, 45, 82)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run("A Benchmarking Study of the Memory Hierarchy")
    sr.font.name = "Times New Roman"; sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x33, 0x55, 0x77)

    add_spacer(doc, 2)

    for line in [
        "Computer Systems | Mid-term Project Report",
        "Istanbul Atlas University",
        "Faculty of Engineering and Natural Sciences",
        "",
        "Team 11",
        "Ebru Inci",
        "Dhafer Hamza Sfaxi",
        "Kasra Nikdel",
        "Fadi Ibrahim Basha",
        "",
        "April 2026",
    ]:
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(line)
        lr.font.name = "Times New Roman"
        lr.font.size = Pt(12 if line else 6)
        lr.font.color.rgb = RGBColor(0x33, 0x33, 0x44)

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ════════════════════════════════════════════════════
    add_heading(doc, "TABLE OF CONTENTS", 1)
    toc_entries = [
        ("1", "Introduction", "3"),
        ("2", "Methodology", "4"),
        ("2.1", "Timer Infrastructure", "4"),
        ("2.2", "Benchmark Architecture", "5"),
        ("2.3", "Platform Notes — Apple Silicon", "5"),
        ("3", "Experiment 1: Stride Access Pattern", "6"),
        ("4", "Experiment 2: Working Set Size", "7"),
        ("5", "Experiment 3: Matrix Traversal Order", "8"),
        ("6", "Experiment 4: Pointer Chasing", "9"),
        ("7", "Experiment 5: Memory Bandwidth", "10"),
        ("8", "Results Dashboard", "12"),
        ("9", "Discussion and Main Findings", "13"),
        ("10", "Conclusions", "14"),
        ("", "References", "15"),
    ]
    for num, title, page in toc_entries:
        tp = doc.add_paragraph()
        tp.paragraph_format.space_after = Pt(2)
        t_run = tp.add_run(f"{num}{'  ' if num else ''}{title}")
        t_run.font.name = "Times New Roman"; t_run.font.size = Pt(11)
        t_run.font.color.rgb = RGBColor(0x11, 0x22, 0x44)
        tab_run = tp.add_run(f"\t{page}")
        tab_run.font.name = "Times New Roman"; tab_run.font.size = Pt(11)
        tab_run.font.color.rgb = RGBColor(0x44, 0x55, 0x66)
    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ════════════════════════════════════════════════════
    add_heading(doc, "1  INTRODUCTION", 1)
    add_body(doc,
        "One of the most consistent surprises for students learning systems programming is just how much "
        "the physical location of data — not the algorithm itself — can dictate program performance. "
        "Modern processors execute instructions at gigahertz frequencies, yet they routinely stall for "
        "tens of nanoseconds waiting for data to arrive from main memory. The only practical remedy is to "
        "keep data close: in registers, then L1 cache, then L2, then L3. Understanding this hierarchy is "
        "not merely academic; it is the difference between code that runs in one second and code that runs "
        "in ten.")
    add_body(doc,
        "This report documents five micro-benchmark experiments carried out as part of the Computer Systems "
        "mid-term project at Istanbul Atlas University. Each experiment targets a specific aspect of the "
        "memory hierarchy: cache-line granularity, effective cache capacities, spatial locality, "
        "random-access latency, and sustained memory bandwidth. All benchmarks were implemented in C++17, "
        "compiled with GCC at optimisation level -O2, and run on an Apple Silicon M-series system.")
    add_body(doc,
        "The Apple Silicon platform introduces important architectural differences from the x86 systems "
        "described in standard textbooks. A Unified Memory Architecture (UMA) eliminates the traditional "
        "off-package DRAM penalty, and a large System Level Cache (SLC) absorbs many sequential access "
        "patterns that would expose latency steps on x86. Where our results deviate from textbook "
        "predictions, we provide hardware-motivated explanations rather than treating them as errors.")
    add_body(doc, "The remainder of this report is structured as follows. Section 2 describes the measurement "
        "methodology and the common timing infrastructure. Sections 3 through 7 present each experiment in "
        "turn, including motivation, implementation, and results. Section 8 provides a unified dashboard "
        "view. Section 9 synthesises the findings, and Section 10 concludes.")
    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # 2. METHODOLOGY
    # ════════════════════════════════════════════════════
    add_heading(doc, "2  METHODOLOGY", 1)

    add_heading(doc, "2.1  Timer Infrastructure", 2)
    add_body(doc,
        "Every experiment shares a single timing primitive defined in src/timer.h. The function "
        "measure_ns<Fn>() accepts a callable (typically a C++ lambda containing the benchmark kernel), "
        "the number of memory accesses that the kernel performs, and an optional iteration count. "
        "It records wall-clock time using std::chrono::high_resolution_clock, which provides "
        "nanosecond resolution without requiring platform-specific POSIX or Windows APIs.")
    add_body(doc,
        "The function runs the callable a fixed number of times — five by default — and returns the "
        "best (minimum) elapsed time divided by the access count. Choosing the minimum rather than "
        "the mean removes the effect of OS scheduling preemptions, page-fault stalls, and thermal "
        "throttling events, all of which introduce upward outliers. The caller supplies the exact "
        "access count so that division is accurate even when the compiler unrolls the inner loop.")
    add_code_block(doc,
        "template<typename Fn>\n"
        "double measure_ns(Fn&& f, long long n_accesses, int iterations = 5) {\n"
        "    using clk = std::chrono::high_resolution_clock;\n"
        "    double best = 1e18;\n"
        "    for (int i = 0; i < iterations; ++i) {\n"
        "        auto t0 = clk::now();\n"
        "        f();\n"
        "        auto t1 = clk::now();\n"
        "        double ns = (double)duration_cast<nanoseconds>(t1 - t0).count()\n"
        "                    / (double)n_accesses;\n"
        "        if (ns < best) best = ns;\n"
        "    }\n"
        "    return best;\n"
        "}"
    )
    add_caption(doc, "Programme 2.1  Core timing template from src/timer.h")

    add_heading(doc, "2.2  Benchmark Architecture", 2)
    add_body(doc,
        "Each experiment follows the same three-phase structure. First, a warm-up pass touches all "
        "pages in the working set, ensuring that virtual memory mappings are established and that "
        "the data is present in at least the L3 cache before timing begins. Second, the timed phase "
        "executes the benchmark kernel inside a measure_ns call. Third, results are accumulated "
        "in a volatile sink variable to prevent the compiler from treating the loop as dead code "
        "and eliminating it entirely.")
    add_body(doc,
        "All five experiments write their results to CSV files in the results/ directory. "
        "The compiler flags -O2 -std=c++17 -Wall -Wextra are fixed for all translation units. "
        "-O2 enables speed optimisations without aggressive auto-vectorisation, which would mask "
        "the memory-hierarchy effects we wish to measure.")

    add_heading(doc, "2.3  Platform Notes — Apple Silicon", 2)
    add_body(doc,
        "The benchmarks were run on an Apple M-series chip. Three architectural properties cause "
        "our results to differ systematically from x86 textbook predictions. The Unified Memory "
        "Architecture places CPU, GPU, and RAM on the same SoC package with LPDDR bandwidth "
        "exceeding 100 GB/s, compared to roughly 50 GB/s for DDR4 DRAM on a typical x86 workstation. "
        "A large System Level Cache, estimated at 8–16 MB depending on the chip variant, sits between "
        "the per-core L2 caches and the LPDDR memory and is shared across all CPU and GPU cores. "
        "Finally, the hardware prefetcher is among the most capable in any consumer processor, "
        "handling sequential and regular-stride patterns that would expose latency on x86.")

    add_data_table(doc,
        ["Property", "Value"],
        [
            ["L1 data cache", "32 KB per P-core"],
            ["L2 cache", "256 KB per P-core"],
            ["L3 / System Level Cache", "~6–16 MB (shared)"],
            ["BENCH_ITERS (best-of-N)", "5"],
            ["MIN_ACCESSES (minimum loop count)", "4,000,000"],
            ["Compiler", "g++ -O2 -std=c++17"],
        ],
        caption="Table 2.1  Hardware and software configuration"
    )
    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # 3. STRIDE
    # ════════════════════════════════════════════════════
    add_heading(doc, "3  EXPERIMENT 1: STRIDE ACCESS PATTERN", 1)
    add_body(doc,
        "The stride experiment probes the granularity of the cache line. When an integer element "
        "is read from a 64 MB array, the hardware does not fetch just four bytes; it fetches the "
        "entire 64-byte cache line that contains the element. If the next accessed element lies "
        "within the same line, the load is free. If it lies in a different line, a new line must "
        "be fetched. Sweeping the stride from 1 element (4 bytes) to 1024 elements (4096 bytes) "
        "therefore reveals the exact byte granularity at which latency transitions from cache-hit "
        "to cache-miss behaviour.")
    add_body(doc,
        "The benchmark allocates a 64 MB array of integers and accesses every stride-th element "
        "in a tight loop. The access count is N / stride so that every run performs the same "
        "amount of arithmetic work and timing differences are purely attributable to memory latency. "
        "A volatile accumulator prevents dead-code elimination.")

    add_figure(doc, figs["stride"], width_in=6.4,
               caption="Figure 3.1  Stride access latency and raw measurements. "
                       "Bars coloured teal indicate L1/L2 cache hits; red indicates cache misses.")

    add_data_table(doc,
        ["Stride (bytes)", "ns/access", "Cache behaviour"],
        [
            ["4", "0.884", "L1 hit — prefetcher absorbs"],
            ["8", "0.887", "L1 hit — prefetcher absorbs"],
            ["16", "0.875", "L1 hit — prefetcher absorbs"],
            ["32", "0.872", "L1 hit — prefetcher absorbs"],
            ["64", "0.928", "L1 hit — one line per access"],
            ["128 ★", "4.500", "MISS — prefetcher saturated"],
            ["256", "3.762", "L2/L3 stream recovers"],
            ["512", "2.419", "L3 streamed"],
            ["1024", "1.610", "L3 streamed"],
            ["2048", "2.295", "L3 / DRAM boundary"],
            ["4096", "2.273", "DRAM prefetch stream"],
        ],
        caption="Table 3.1  Full stride results. ★ marks the first significant latency spike."
    )
    add_body(doc,
        "The latency remains flat at approximately 0.88 ns for strides up to 64 bytes, "
        "then rises sharply to 4.5 ns at 128 bytes — a 5× penalty. On x86, textbooks typically "
        "place this transition at the 64-byte cache-line boundary itself. On Apple Silicon, the "
        "hardware prefetcher successfully hides the one-line stride, so the penalty only appears "
        "when two consecutive accesses touch two consecutive cache lines (stride = 128 bytes). "
        "After the spike, latency gradually recovers as the access pattern becomes a predictable "
        "streaming read that the L2 and L3 prefetchers can partially service.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # 4. WORKING SET
    # ════════════════════════════════════════════════════
    add_heading(doc, "4  EXPERIMENT 2: WORKING SET SIZE", 1)
    add_body(doc,
        "Where the stride experiment varies the spatial regularity of access, the working-set "
        "experiment varies the total data footprint. The access stride is fixed at 1 (sequential), "
        "and the array size is swept from 1 KB to 256 MB in powers of two. As the array grows "
        "beyond each cache level's capacity, the effective latency should step upward: L1 → L2 "
        "→ L3 → DRAM, producing three distinct transition points and four latency plateaus.")
    add_body(doc,
        "To ensure that even tiny arrays accumulate enough elapsed time for accurate timing, "
        "small arrays are scanned repeatedly. The total number of accesses is kept at a minimum "
        "of 4,000,000 by scaling the number of passes with MIN_ACCESSES / n.")

    add_figure(doc, figs["workingset"], width_in=6.2,
               caption="Figure 4.1  Working-set latency sweep. Coloured bands denote the theoretical "
                       "L1, L2, L3, and DRAM regions. The dashed amber line marks ~1 ns.")

    add_body(doc,
        "The results show a completely flat curve at approximately 1 ns across all array sizes "
        "from 1 KB to 256 MB. No transitions are visible. This outcome, while initially surprising, "
        "is fully explained by the Apple Silicon architecture. The hardware prefetcher recognises "
        "the sequential scan pattern and issues prefetch requests far enough in advance that the "
        "data arrives before the CPU stalls. The System Level Cache absorbs repeated scans of "
        "arrays up to approximately 16 MB, and the high-bandwidth LPDDR memory services larger "
        "arrays faster than DDR4-based systems would. The textbook four-step latency profile is "
        "an x86 phenomenon on systems with weaker prefetchers and higher DRAM latency.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # 5. MATRIX
    # ════════════════════════════════════════════════════
    add_heading(doc, "5  EXPERIMENT 3: MATRIX TRAVERSAL ORDER", 1)
    add_body(doc,
        "C and C++ store two-dimensional arrays in row-major order: the elements of row i occupy "
        "a contiguous block of memory, immediately followed by the elements of row i+1. A "
        "row-major traversal — incrementing the column index in the inner loop — therefore "
        "accesses addresses in ascending sequential order, which is optimal for the prefetcher "
        "and for cache-line utilisation. A column-major traversal — incrementing the row index "
        "in the inner loop — accesses addresses separated by N × 8 bytes (the row stride), "
        "which is highly non-sequential for large N.")
    add_body(doc,
        "The theoretical maximum slowdown is 8× for very large matrices. A 64-byte cache line "
        "holds exactly eight consecutive doubles. Row-major access uses all eight; column-major "
        "access uses only one, wasting the remaining seven. In practice, modern hardware partially "
        "compensates through prefetching.")

    add_figure(doc, figs["matrix"], width_in=6.4,
               caption="Figure 5.1  Left: absolute traversal latency (ns/element). "
                       "Right: column-major slowdown ratio. Amber bars exceed the 1.3× significance threshold. "
                       "The dashed red line marks the theoretical 8× limit.")

    add_data_table(doc,
        ["N", "Matrix size", "Row-major (ns)", "Col-major (ns)", "Slowdown"],
        [
            ["64",   "32 KB",   "1.810", "1.965", "1.09×"],
            ["128",  "128 KB",  "1.889", "2.248", "1.19×"],
            ["256",  "512 KB",  "1.928", "2.832", "1.47×"],
            ["512",  "2 MB",    "1.949", "2.836", "1.45×"],
            ["1024", "8 MB",    "1.959", "2.840", "1.45×"],
            ["2048", "32 MB",   "1.966", "2.923", "1.49×"],
            ["4096", "128 MB",  "1.967", "3.004", "1.53×"],
        ],
        caption="Table 5.1  Matrix traversal results"
    )
    add_body(doc,
        "The penalty is negligible for small matrices (N ≤ 128) that fit entirely within L1 or L2 "
        "cache — both traversal orders keep the data hot, so spatial locality is irrelevant. "
        "Once the matrix exceeds L2 capacity (N = 256, 512 KB), the penalty stabilises at "
        "approximately 1.45–1.53×. The gap between observed (~1.5×) and theoretical (8×) is "
        "attributable to the hardware prefetcher, which recognises the constant column stride "
        "and partially services it. Column-major traversal remains systematically and measurably "
        "slower; it is not benign.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # 6. LINKED LIST
    # ════════════════════════════════════════════════════
    add_heading(doc, "6  EXPERIMENT 4: POINTER CHASING", 1)
    add_body(doc,
        "The pointer-chasing experiment constructs a singly-linked list whose nodes are arranged "
        "in a uniformly shuffled random order. Each node occupies exactly one 64-byte cache line, "
        "enforced by the alignas(64) attribute and a 56-byte padding array. After allocating N "
        "nodes, a Fisher-Yates shuffle produces a random permutation of node pointers, and the "
        "benchmark traverses the list by following next pointers. Because each successive access "
        "address is unknown until the previous load completes, the CPU cannot issue speculative "
        "prefetch requests and must wait for the full memory latency at every hop.")
    add_body(doc,
        "This creates a load-use dependency chain that completely serialises the traversal: hop i+1 "
        "cannot begin until hop i's cache miss has been resolved. The effective latency measured "
        "therefore reflects the round-trip time to whichever cache level holds the target node.")
    add_body(doc,
        "A critical implementation detail was required to obtain correct measurements on Apple Silicon: "
        "the traversal must be placed in a __attribute__((noinline)) helper function, and each pointer "
        "dereference must be followed by an asm volatile(\"\" ::: \"memory\") compiler barrier. "
        "Without these, the compiler at -O2 collapses the dependency chain and the measured time is "
        "near zero. The asm volatile barrier emits no CPU instructions but prevents the compiler "
        "from reordering or eliminating any memory access across the boundary.")

    add_figure(doc, figs["linkedlist"], width_in=6.2,
               caption="Figure 6.1  Pointer-chasing latency vs list size. "
                       "Shaded bands show expected cache-level regions. "
                       "The clear step from ~0.65 ns to ~83 ns demonstrates the prefetcher's defeat.")

    add_data_table(doc,
        ["List size", "ns/hop", "Cache level"],
        [
            ["4 KB – 128 KB",   "~0.65 ns",   "L1/L2 — data fits in fast cache"],
            ["256 KB – 512 KB", "~4–11 ns",   "L2/L3 boundary transition"],
            ["1 MB – 16 MB",    "~8–9 ns",    "L3 cache"],
            ["32 MB",           "~38 ns",     "L3/DRAM boundary"],
            ["64 MB – 256 MB",  "~70–83 ns",  "DRAM — full memory round-trip"],
        ],
        caption="Table 6.1  Pointer-chasing latency summary by cache level"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # 7. BANDWIDTH
    # ════════════════════════════════════════════════════
    add_heading(doc, "7  EXPERIMENT 5: MEMORY BANDWIDTH", 1)
    add_body(doc,
        "Two streaming kernels are applied to double-precision floating-point arrays swept from "
        "4 KB to 256 MB. The copy kernel writes every element of a destination array from a "
        "corresponding source element (dst[i] = src[i]), performing one read and one write per "
        "element. The sum kernel accumulates all array elements into a scalar (s += src[i]), "
        "performing one read per element. Both kernels iterate over a contiguous allocation, "
        "which is the best possible access pattern for the prefetcher.")
    add_body(doc,
        "Bandwidth is derived from the measured nanoseconds per access. For the copy kernel, "
        "each access transfers sizeof(double) = 8 bytes and involves both a read and a write, "
        "so the bandwidth is 2 × 8 bytes / (ns × 10⁻⁹ seconds) / 10⁹ = GB/s. For the sum "
        "kernel, only the read side contributes, so the multiplier is 1.")
    add_code_block(doc,
        "// ns/access -> GB/s conversion:\n"
        "double copy_gb = (sizeof(double) * 2.0) / (copy_ns * 1e-9) / 1e9;  // read + write\n"
        "double sum_gb  =  sizeof(double)        / (sum_ns  * 1e-9) / 1e9;  // read only"
    )
    add_caption(doc, "Programme 7.1  GB/s conversion formula from src/bandwidth_bench.cpp")

    add_figure(doc, figs["bandwidth"], width_in=6.4,
               caption="Figure 7.1  Left: full bandwidth vs array-size sweep. "
                       "Right: peak bandwidth grouped by cache level. "
                       "Note that sum bandwidth is flat across all levels despite varying copy bandwidth.")

    add_data_table(doc,
        ["Array size", "Copy (GB/s)", "Sum (GB/s)", "Cache level"],
        [
            ["4 KB",   "249.9",  "17.2",  "L1"],
            ["8 KB",   "253.0",  "17.2",  "L1 (peak copy)"],
            ["32 KB",  "168.1",  "17.2",  "L1/L2 boundary"],
            ["256 KB", "170.3",  "17.2",  "L2"],
            ["4 MB",   "167.1",  "17.2",  "L2/L3"],
            ["8 MB",   "88.7",   "17.2",  "L3"],
            ["64 MB",  "109.9",  "17.2",  "SLC/DRAM"],
            ["256 MB", "109.3",  "17.1",  "DRAM"],
        ],
        caption="Table 7.1  Selected bandwidth measurements"
    )
    add_body(doc,
        "Copy bandwidth peaks at 253 GB/s for small arrays fully resident in L1 and falls to "
        "approximately 109 GB/s for DRAM-sized arrays — still a remarkable figure that reflects "
        "Apple Silicon's high-bandwidth LPDDR. Sum bandwidth, by contrast, is completely flat "
        "at approximately 17.2 GB/s regardless of array size. This is not a memory effect. "
        "The sum kernel contains a serial floating-point add dependency: each iteration "
        "s += src[i] must wait for the previous addition to complete before it can begin. "
        "With a floating-point add latency of roughly five cycles at 3.5 GHz, throughput is "
        "theoretically capped near 5.6 GB/s; the observed 17 GB/s suggests mild out-of-order "
        "overlap but confirms that the bottleneck is arithmetic latency, not memory bandwidth.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # 8. DASHBOARD
    # ════════════════════════════════════════════════════
    add_heading(doc, "8  RESULTS DASHBOARD", 1)
    add_body(doc,
        "Figure 8.1 presents all four primary experiments in a single dashboard view, allowing "
        "direct visual comparison of the different memory-hierarchy effects. The top-left panel "
        "shows the stride latency spike; the top-right panel shows the flat working-set curve "
        "characteristic of Apple Silicon; the bottom-left panel shows matrix traversal slowdown "
        "ratios; and the bottom-right panel shows the bandwidth divergence between copy and sum kernels.")
    add_figure(doc, figs["dashboard"], width_in=6.8,
               caption="Figure 8.1  Unified results dashboard. Each panel corresponds to one experiment. "
                       "Top-left: stride. Top-right: working set. Bottom-left: matrix slowdown. "
                       "Bottom-right: memory bandwidth.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # 9. DISCUSSION
    # ════════════════════════════════════════════════════
    add_heading(doc, "9  DISCUSSION AND MAIN FINDINGS", 1)
    add_body(doc,
        "Taken together, the five experiments paint a consistent picture: memory access pattern "
        "is a first-order performance variable, independent of algorithmic complexity. Two programs "
        "with identical asymptotic complexity can differ in runtime by an order of magnitude if "
        "their data access patterns differ.")

    add_heading(doc, "9.1  Cache-Line Granularity (Experiment 1)", 2)
    add_body(doc,
        "The 64-byte cache line is the atomic unit of data transfer between DRAM and cache. "
        "Accessing elements spaced more than one cache line apart wastes the prefetched data "
        "on each line and forces a new fetch for every element. On Apple Silicon the hardware "
        "prefetcher extends the effective granularity to 128 bytes, but the underlying principle "
        "holds: spatial locality within 64 bytes is free; spatial locality beyond it is costly.")

    add_heading(doc, "9.2  Working Set and Cache Capacity (Experiment 2)", 2)
    add_body(doc,
        "Sequential access is the best-case scenario for the hardware prefetcher. On Apple Silicon, "
        "the combination of an aggressive prefetcher, the System Level Cache, and high-bandwidth LPDDR "
        "renders the cache hierarchy nearly invisible for sequential workloads. On x86 hardware the "
        "four-tier latency profile would be clearly observable. This experiment underscores that "
        "platform awareness matters: results that look surprising on one architecture are often "
        "fully explicable on another.")

    add_heading(doc, "9.3  Spatial Locality and Matrix Traversal (Experiment 3)", 2)
    add_body(doc,
        "Column-major traversal of a row-major C array remains consistently and measurably slower "
        "than row-major traversal for all matrices larger than the L2 cache. The measured 1.5× "
        "slowdown on Apple Silicon would likely be 4–8× on a typical x86 system with a weaker "
        "prefetcher. The practical lesson is unchanged: write loops that access memory in "
        "continuous, predictable order.")

    add_heading(doc, "9.4  Pointer Chasing and Prefetcher Defeat (Experiment 4)", 2)
    add_body(doc,
        "The pointer-chasing experiment provides the clearest demonstration of the cost of "
        "irregular memory access. Latency rises from 0.65 ns when the list fits in L1/L2 to "
        "83 ns when it is in DRAM — a 128× difference. This is not a workload-size effect; it "
        "is a pure latency effect caused by the serialised load-use dependency chain. No "
        "prefetcher can help because the next address is not known until the current load completes. "
        "The experiment also required an implementation workaround (noinline + asm volatile barrier) "
        "that illustrates a further practical concern: modern compilers and CPUs can optimise away "
        "benchmarks that are not carefully protected against dead-code elimination.")

    add_heading(doc, "9.5  Memory Bandwidth and Serial Dependencies (Experiment 5)", 2)
    add_body(doc,
        "Peak copy bandwidth of 253 GB/s demonstrates Apple Silicon's raw memory throughput "
        "advantage. The flat 17 GB/s sum bandwidth demonstrates that high memory bandwidth is "
        "not always the binding constraint. Any algorithm featuring a serial carried dependency — "
        "such as prefix-sum, running maximum, or leaky integrator — will be bottlenecked by "
        "compute latency regardless of how fast the memory subsystem can supply data. Breaking "
        "such dependencies through SIMD intrinsics or algorithm restructuring is one of the most "
        "high-leverage optimisations available.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # 10. CONCLUSIONS
    # ════════════════════════════════════════════════════
    add_heading(doc, "10  CONCLUSIONS", 1)
    add_body(doc,
        "This project set out to make the memory hierarchy tangible through direct measurement, "
        "and the experiments fulfilled that goal. Five distinct hardware phenomena were isolated "
        "and quantified with sub-nanosecond precision using only standard C++17 and the "
        "chrono library. The key findings are summarised below.")

    findings = [
        ("Cache-line granularity is real and measurable.",
         "A 5× latency increase appears at the 128-byte stride boundary on Apple Silicon, "
         "confirming that the hardware prefetcher has finite reach. On x86, the boundary is at 64 bytes."),
        ("Sequential access is extremely forgiving on Apple Silicon.",
         "The working-set experiment found no visible latency transitions, owing to the hardware "
         "prefetcher, the System Level Cache, and high-bandwidth LPDDR. Sequential scans are "
         "effectively immune to the memory hierarchy on this platform."),
        ("Traversal order matters even when the prefetcher compensates.",
         "Column-major matrix access is 1.5× slower than row-major for large matrices, and this "
         "penalty would be 4–8× on x86. Writing cache-friendly loops is always worthwhile."),
        ("Irregular access enforces full memory latency.",
         "Pointer chasing on a shuffled linked list raises per-element latency from 0.65 ns (L1) "
         "to 83 ns (DRAM), a 128× range. Contiguous data structures are far preferable for "
         "iteration-heavy workloads."),
        ("Memory bandwidth is not always the bottleneck.",
         "The sum kernel plateaus at 17 GB/s regardless of array size, bounded by a serial "
         "floating-point dependency chain. High memory bandwidth only helps when the algorithm "
         "structure can actually exploit it."),
    ]

    for i, (title, body) in enumerate(findings, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pr = p.add_run(f"Finding {i}: {title}  ")
        pr.bold = True; pr.font.size = Pt(12)
        pr.font.name = "Times New Roman"
        pr.font.color.rgb = RGBColor(11, 45, 82)
        br = p.add_run(body)
        br.font.size = Pt(12); br.font.name = "Times New Roman"
        br.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        p.paragraph_format.line_spacing = Pt(18)
        p.paragraph_format.space_after  = Pt(8)

    add_spacer(doc)
    add_body(doc,
        "The broader lesson is that writing performant systems code requires understanding the hardware "
        "it runs on. Algorithm complexity gives a platform-independent upper bound on performance, but "
        "the constant factors — determined almost entirely by memory access patterns — often dominate "
        "in practice. The ability to measure, interpret, and act on memory-hierarchy data is one of the "
        "most valuable skills a systems programmer can cultivate.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # REFERENCES
    # ════════════════════════════════════════════════════
    add_heading(doc, "REFERENCES", 1)
    refs = [
        "R. E. Bryant and D. R. O'Hallaron, Computer Systems: A Programmer's Perspective, "
        "3rd ed. Pearson, 2016.",
        "U. Drepper, 'What Every Programmer Should Know About Memory,' Red Hat, Inc., 2007. "
        "[Online]. Available: https://www.akkadia.org/drepper/cpumemory.pdf",
        "Intel Corporation, Intel® 64 and IA-32 Architectures Optimization Reference Manual, "
        "Order No. 248966, 2023.",
        "W. Stallings, Operating Systems: Internals and Design Principles, 9th ed. Pearson, 2018.",
        "Apple Inc., 'Apple Silicon Overview,' Apple Developer Documentation, 2023. "
        "[Online]. Available: https://developer.apple.com/documentation/apple-silicon",
        "ISO/IEC, Programming Languages — C++, ISO/IEC 14882:2017 (C++17), 2017.",
    ]
    for i, ref in enumerate(refs, 1):
        rp = doc.add_paragraph()
        rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rr = rp.add_run(f"[{i}]  {ref}")
        rr.font.name  = "Times New Roman"
        rr.font.size  = Pt(11)
        rr.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        rp.paragraph_format.space_after      = Pt(6)
        rp.paragraph_format.left_indent      = Cm(0.7)
        rp.paragraph_format.first_line_indent = Cm(-0.7)

    return doc


# ═══════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("\n── Generating charts ─────────────────────────────────────────")
    figs = {
        "stride":     fig_stride(),
        "workingset": fig_workingset(),
        "matrix":     fig_matrix(),
        "linkedlist": fig_linkedlist(),
        "bandwidth":  fig_bandwidth(),
    }
    stride_data = load_csv("stride.csv")
    ws_data     = load_csv("workingset.csv")
    matrix_data = load_csv("matrix.csv")
    bw_data     = load_csv("bandwidth.csv")
    figs["dashboard"] = fig_dashboard(stride_data, ws_data, matrix_data, bw_data)

    print("\n── Assembling DOCX ───────────────────────────────────────────")
    doc = build_doc(figs)
    doc.save(str(OUT))
    print(f"\n✅  Report saved to: {OUT}")
    print(f"    Size: {OUT.stat().st_size // 1024} KB")
